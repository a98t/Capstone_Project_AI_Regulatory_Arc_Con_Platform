"""
Document parser — extracts structured text from regulatory documents.

Supports:
- Text-based PDFs (via PyMuPDF)
- Scanned PDFs (OCR fallback via pytesseract, if installed)
- Word documents .docx (via python-docx)
- Plain text .txt files

Returns a ParsedDocument with full text and extracted metadata.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

import structlog

log = structlog.get_logger(__name__)


@dataclass
class ParsedDocument:
    file_path: str
    doc_hash: str
    title: str
    doc_number: str          # e.g. "СП РК 2.02-101-2015"
    doc_type: str            # e.g. "СНиП", "СП", "ҚНжЕ", "СТ РК"
    year: Optional[int]
    language: str            # "ru", "kk", "en"
    full_text: str
    page_count: int
    extraction_method: str   # "pymupdf" or "ocr"


# Regex patterns to detect document type from title/filename
_DOC_TYPE_PATTERNS = [
    (r"СНиП", "СНиП"),
    (r"СП\s+РК", "СП РК"),
    (r"СП\s+EN", "СП EN"),
    (r"ҚНжЕ", "ҚНжЕ"),
    (r"СТ\s+РК", "СТ РК"),
    (r"РДС\s+РК", "РДС РК"),
    (r"ВСН", "ВСН"),
    (r"ПР\s+РК", "ПР РК"),
    (r"СН\s+РК", "СН РК"),
]

_YEAR_PATTERN = re.compile(r"\b(19[89]\d|20[012]\d)\b")


def _detect_doc_type(text: str) -> str:
    for pattern, doc_type in _DOC_TYPE_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            return doc_type
    return "UNKNOWN"


def _detect_year(text: str) -> Optional[int]:
    match = _YEAR_PATTERN.search(text)
    return int(match.group()) if match else None


def _detect_language(text: str) -> str:
    """Simple heuristic: count Cyrillic vs Latin characters."""
    cyrillic = len(re.findall(r"[а-яёА-ЯЁ]", text))
    kazakh = len(re.findall(r"[әіңғүұқөһӘІҢҒҮҰҚӨҺ]", text))
    latin = len(re.findall(r"[a-zA-Z]", text))
    if kazakh > 50:
        return "kk"
    if cyrillic > latin:
        return "ru"
    return "en"


def _compute_hash(path: Path) -> str:
    sha256 = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _extract_with_pymupdf(path: Path) -> tuple[str, int]:
    """Extract text from all pages. Returns (full_text, page_count)."""
    doc = fitz.open(str(path))
    pages_text = []
    for page in doc:
        pages_text.append(page.get_text("text"))
    doc.close()
    return "\n".join(pages_text), len(pages_text)


def _extract_with_ocr(path: Path) -> tuple[str, int]:
    """OCR fallback for scanned PDFs. Requires pytesseract + poppler."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(str(path), dpi=200)
        pages_text = [
            pytesseract.image_to_string(img, lang="rus+kaz+eng")
            for img in images
        ]
        return "\n".join(pages_text), len(pages_text)
    except ImportError:
        log.warning("ocr_unavailable", path=str(path), hint="Install pytesseract and pdf2image for OCR support")
        return "", 0


def parse_pdf(path: Path) -> Optional[ParsedDocument]:
    """
    Parse a single PDF file.

    Returns None if the file cannot be parsed (logged as warning).
    """
    path = Path(path)
    if not path.exists() or path.suffix.lower() != ".pdf":
        log.warning("invalid_file", path=str(path))
        return None

    log.info("parsing_pdf", path=str(path))
    doc_hash = _compute_hash(path)

    full_text, page_count = _extract_with_pymupdf(path)
    extraction_method = "pymupdf"

    # If extraction produced almost no text → likely scanned → try OCR
    if len(full_text.strip()) < 100:
        log.info("low_text_fallback_ocr", path=str(path), char_count=len(full_text))
        full_text, page_count = _extract_with_ocr(path)
        extraction_method = "ocr"

    if len(full_text.strip()) < 50:
        log.warning("parse_failed", path=str(path), reason="Insufficient text extracted")
        return None

    # Extract metadata from first 500 characters (document header area)
    header = full_text[:500]
    stem = path.stem  # filename without extension

    title = stem.replace("_", " ")
    doc_number = stem  # Will be refined by chunker from document content
    doc_type = _detect_doc_type(header + stem)
    year = _detect_year(header + stem)
    language = _detect_language(full_text[:2000])

    return ParsedDocument(
        file_path=str(path),
        doc_hash=doc_hash,
        title=title,
        doc_number=doc_number,
        doc_type=doc_type,
        year=year,
        language=language,
        full_text=full_text,
        page_count=page_count,
        extraction_method=extraction_method,
    )


def parse_docx(path: Path) -> Optional["ParsedDocument"]:
    """
    Parse a Microsoft Word (.docx) regulatory document.
    Extracts text from paragraphs and tables.
    """
    path = Path(path)
    if not path.exists() or path.suffix.lower() != ".docx":
        log.warning("invalid_docx_file", path=str(path))
        return None

    log.info("parsing_docx", path=str(path))
    doc_hash = _compute_hash(path)

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
    except Exception as exc:
        log.warning("docx_open_failed", path=str(path), error=str(exc))
        return None

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cells (regulations often use tables for norm values)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts)

    if len(full_text.strip()) < 50:
        log.warning("parse_failed", path=str(path), reason="Insufficient text in docx")
        return None

    header = full_text[:500]
    stem = path.stem
    title = stem.replace("_", " ")
    doc_number = stem
    doc_type = _detect_doc_type(header + stem)
    year = _detect_year(header + stem)
    language = _detect_language(full_text[:2000])

    return ParsedDocument(
        file_path=str(path),
        doc_hash=doc_hash,
        title=title,
        doc_number=doc_number,
        doc_type=doc_type,
        year=year,
        language=language,
        full_text=full_text,
        page_count=len(doc.paragraphs) // 30 + 1,  # rough estimate
        extraction_method="docx",
    )


def parse_text_file(path: Path) -> Optional["ParsedDocument"]:
    """
    Parse a plain-text (.txt) regulatory document.
    Useful for seed/demo data before real PDFs are available.
    """
    path = Path(path)
    if not path.exists() or path.suffix.lower() != ".txt":
        log.warning("invalid_text_file", path=str(path))
        return None

    log.info("parsing_text", path=str(path))
    doc_hash = _compute_hash(path)

    full_text = path.read_text(encoding="utf-8", errors="replace")

    if len(full_text.strip()) < 50:
        log.warning("parse_failed", path=str(path), reason="Insufficient text")
        return None

    header = full_text[:500]
    stem = path.stem
    title = stem.replace("_", " ")
    doc_number = stem
    doc_type = _detect_doc_type(header + stem)
    year = _detect_year(header + stem)
    language = _detect_language(full_text[:2000])

    return ParsedDocument(
        file_path=str(path),
        doc_hash=doc_hash,
        title=title,
        doc_number=doc_number,
        doc_type=doc_type,
        year=year,
        language=language,
        full_text=full_text,
        page_count=1,
        extraction_method="text",
    )
